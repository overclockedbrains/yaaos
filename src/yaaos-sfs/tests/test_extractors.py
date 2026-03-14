"""Tests for the extractor registry and built-in extractors."""

from __future__ import annotations

import pytest
from pathlib import Path

from yaaos_sfs.extractors import extract_text, get_extractor, get_supported_extensions
from yaaos_sfs.extractors.text import extract_plaintext


class TestExtractorRegistry:
    def test_text_extensions_registered(self):
        """Core text extensions should always be registered."""
        exts = get_supported_extensions()
        for ext in [".py", ".js", ".ts", ".md", ".txt", ".json", ".yaml", ".toml"]:
            assert ext in exts, f"{ext} should be registered"

    def test_pdf_registered(self):
        assert get_extractor(Path("test.pdf")) is not None

    def test_unknown_extension_returns_none(self):
        assert get_extractor(Path("test.xyz_unknown")) is None

    def test_extract_text_unknown_returns_none(self, tmp_path):
        f = tmp_path / "test.xyz_unknown_ext"
        f.write_text("hello")
        assert extract_text(f) is None

    def test_extract_text_nonexistent_file(self, tmp_path):
        f = tmp_path / "nope.py"
        result = extract_text(f)
        assert result is None


class TestPlaintextExtractor:
    def test_extract_python(self, tmp_path):
        f = tmp_path / "hello.py"
        f.write_text("def greet():\n    return 'hi'\n")
        text = extract_plaintext(f)
        assert "def greet" in text

    def test_extract_markdown(self, tmp_path):
        f = tmp_path / "notes.md"
        f.write_text("# Title\n\nSome content here.\n")
        text = extract_plaintext(f)
        assert "# Title" in text
        assert "Some content" in text

    def test_extract_json(self, tmp_path):
        f = tmp_path / "config.json"
        f.write_text('{"key": "value"}')
        text = extract_plaintext(f)
        assert '"key"' in text

    def test_extract_utf8_special(self, tmp_path):
        f = tmp_path / "unicode.txt"
        f.write_text("Héllo 日本語 🎉", encoding="utf-8")
        text = extract_plaintext(f)
        assert "Héllo" in text

    def test_extract_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        text = extract_plaintext(f)
        assert text == ""

    def test_extract_nonexistent(self, tmp_path):
        f = tmp_path / "nope.txt"
        assert extract_plaintext(f) is None


class TestDocumentExtractors:
    """Tests for document extractors — skipped if deps not installed."""

    def test_docx_extraction(self, tmp_path):
        """Test DOCX extraction if python-docx is available."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from yaaos_sfs.extractors.documents import extract_docx

        # Create a simple DOCX
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("This is test content.")
        doc.add_paragraph("Another paragraph.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        text = extract_docx(docx_path)
        assert text is not None
        assert "Test Heading" in text
        assert "test content" in text

    def test_xlsx_extraction(self, tmp_path):
        """Test XLSX extraction if openpyxl is available."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        from yaaos_sfs.extractors.documents import extract_xlsx

        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Age", "City"])
        ws.append(["Alice", 30, "NYC"])
        ws.append(["Bob", 25, "London"])
        xlsx_path = tmp_path / "test.xlsx"
        wb.save(str(xlsx_path))

        text = extract_xlsx(xlsx_path)
        assert text is not None
        assert "Data" in text
        assert "Alice" in text
        assert "Bob" in text

    def test_rtf_extraction(self, tmp_path):
        """Test RTF extraction if striprtf is available."""
        try:
            from striprtf.striprtf import rtf_to_text  # noqa: F401
        except ImportError:
            pytest.skip("striprtf not installed")

        from yaaos_sfs.extractors.documents import extract_rtf

        rtf_content = r"{\rtf1\ansi Hello RTF World!}"
        rtf_path = tmp_path / "test.rtf"
        rtf_path.write_text(rtf_content)

        text = extract_rtf(rtf_path)
        assert text is not None
        assert "Hello" in text

    def test_corrupt_docx_returns_none(self, tmp_path):
        """Corrupt files should return None, not crash."""
        try:
            import docx  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        from yaaos_sfs.extractors.documents import extract_docx

        bad = tmp_path / "corrupt.docx"
        bad.write_bytes(b"not a real docx file")
        result = extract_docx(bad)
        assert result is None


class TestMediaExtractors:
    """Tests for media metadata extractors — skipped if deps not installed."""

    def test_image_metadata(self, tmp_path):
        """Test image metadata extraction if Pillow is available."""
        try:
            from PIL import Image
        except ImportError:
            pytest.skip("Pillow not installed")

        from yaaos_sfs.extractors.media import extract_image_metadata

        # Create a simple PNG
        img = Image.new("RGB", (100, 100), color="red")
        img_path = tmp_path / "test.png"
        img.save(str(img_path))

        text = extract_image_metadata(img_path)
        assert text is not None
        assert "Image: test.png" in text
        assert "100x100" in text

    def test_audio_metadata(self, tmp_path):
        """Test audio metadata extraction if mutagen is available."""
        try:
            import mutagen  # noqa: F401
        except ImportError:
            pytest.skip("mutagen not installed")

        from yaaos_sfs.extractors.media import extract_audio_metadata

        # Create a minimal valid MP3 (just test that it doesn't crash on non-audio)
        bad = tmp_path / "fake.mp3"
        bad.write_bytes(b"\x00" * 100)
        result = extract_audio_metadata(bad)
        # Should return None for invalid audio, not crash
        assert result is None

    def test_corrupt_image_returns_none(self, tmp_path):
        """Corrupt images should return None, not crash."""
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            pytest.skip("Pillow not installed")

        from yaaos_sfs.extractors.media import extract_image_metadata

        bad = tmp_path / "corrupt.jpg"
        bad.write_bytes(b"not an image")
        result = extract_image_metadata(bad)
        assert result is None
